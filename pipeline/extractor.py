import os

def extract_text(file_path: str) -> str:
    """
    Extracts plain text from a file.
    Supports: .txt, .pdf, .docx
    Should support .doc in future
    Raises ValueError for unsupported formats.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        return _extract_txt(file_path)
    elif ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: '{ext}'. Supported: .txt, .pdf, .docx")


def extract_text_from_bytes(data: bytes, filename):
    """
    Extracts plain text from raw bytes (e.g. from a multipart upload).
    Uses the filename to determine format.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt":
        return data.decode("utf-8", errors="replace")
    elif ext == ".pdf":
        return _extract_pdf_from_bytes(data)
    elif ext == ".docx":
        return _extract_docx_from_bytes(data)
    else:
        raise ValueError(f"Unsupported file format: '{ext}'. Supported: .txt, .pdf, .docx")


# --- private helpers ---

def _extract_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_pdf(file_path):
    import pdfplumber
    pages = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n".join(pages)


def _extract_pdf_from_bytes(data: bytes):
    import pdfplumber
    import io
    pages = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n".join(pages)


def _extract_docx(file_path):
    from docx import Document
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _extract_docx_from_bytes(data: bytes) -> str:
    from docx import Document
    import io
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())